"""DOCX parser using python-docx."""

from __future__ import annotations

import io
from typing import Any

from docx import Document

from ..schemas import FilePreview
from .base import ParseError, Parser

_PREVIEW_CHARS = 12_000


class DocxParser(Parser):
    kind = "docx"

    def _open(self, data: bytes) -> Document:
        try:
            return Document(io.BytesIO(data))
        except Exception as e:
            raise ParseError(f"Failed to read .docx: {e}") from e

    def _all_text(self, doc: Document) -> str:
        return "\n".join(p.text for p in doc.paragraphs if p.text)

    def meta(self, data: bytes, *, filename: str) -> dict[str, Any]:
        doc = self._open(data)
        return {
            "kind": self.kind,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }

    def preview(self, data: bytes, *, filename: str) -> FilePreview:
        doc = self._open(data)
        text = self._all_text(doc)
        truncated = len(text) > _PREVIEW_CHARS
        if truncated:
            text = text[:_PREVIEW_CHARS] + "\n\n…(truncated)"
        return FilePreview(kind=self.kind, text=text, truncated=truncated)
